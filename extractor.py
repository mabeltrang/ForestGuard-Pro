import os
from pypdf import PdfReader
from docx import Document
import pandas as pd


def _detectar_conteo_inventario(df_dict: dict) -> dict:
    """
    Detecta si alguna hoja del Excel es un inventario forestal (columnas tipo
    ID, Nombre científico, CAP/DAP, Cobertura, etc.) y cuenta el número REAL
    de árboles contando FILAS DE DATOS.

    IMPORTANTE: el conteo nunca se basa en el valor de la columna ID (que
    puede tener saltos, duplicados o no iniciar en 1) — se cuenta la columna
    de Nombre científico/común como ancla, fila por fila, hasta la primera
    fila vacía que indique el fin de la tabla.

    Retorna {nombre_hoja: {"individuos": int, "volumen_m3": float|None}}.
    """
    resultados = {}
    for sheet_name, df in df_dict.items():
        header_idx = None
        for i in range(min(15, len(df))):
            fila = df.iloc[i].astype(str).str.strip().str.lower()
            valores = set(fila.tolist())
            if "id" in valores and any(
                any(k in v for k in ["científico", "cientifico", "cap", "dap", "cobertura"])
                for v in valores
            ):
                header_idx = i
                break
        if header_idx is None:
            continue

        header_row = [str(v).strip() for v in df.iloc[header_idx].tolist()]
        header_row_lower = [v.lower() for v in header_row]
        datos = df.iloc[header_idx + 1:]

        # Columna ancla para contar individuos: nombre científico/común,
        # NUNCA la columna ID (para no depender de sus valores).
        col_ancla = None
        for j, val in enumerate(header_row_lower):
            if "científico" in val or "cientifico" in val or "nombre común" in val or "nombre comun" in val:
                col_ancla = j
                break
        if col_ancla is None:
            col_ancla = 1 if len(header_row) > 1 else 0

        conteo = 0
        for val in datos.iloc[:, col_ancla]:
            if pd.isna(val) or str(val).strip() == "":
                if conteo > 0:
                    break
                continue
            conteo += 1

        if conteo == 0:
            continue

        # Volumen total (VT), si existe la columna, sumado sobre las mismas
        # filas contadas como individuos.
        volumen_m3 = None
        for j, val in enumerate(header_row_lower):
            if val.startswith("vt ") or val == "vt" or "vt (m3)" in val:
                serie_vt = pd.to_numeric(datos.iloc[:conteo, j], errors="coerce")
                volumen_m3 = round(float(serie_vt.sum()), 3)
                break

        resultados[sheet_name] = {"individuos": conteo, "volumen_m3": volumen_m3}

    return resultados


def extract_text_from_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"

        elif ext in [".docx", ".doc"]:
            text = extract_docx(file_path)

        elif ext in [".xlsx", ".xls"]:
            # Detección de inventario forestal: se lee primero sin asumir
            # encabezado (los inventarios suelen traer 3-4 filas de
            # metadatos --proyecto, propietario, ubicación, fecha-- antes de
            # la fila de columnas real) para ubicar la tabla y contar árboles
            # por FILAS DE DATOS, nunca por el valor de la columna ID.
            try:
                df_dict_raw = pd.read_excel(file_path, sheet_name=None, header=None)
                conteos = _detectar_conteo_inventario(df_dict_raw)
            except Exception:
                conteos = {}

            df_dict = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, df in df_dict.items():
                text += f"\n[Hoja: {sheet_name}]\n"
                if sheet_name in conteos:
                    info = conteos[sheet_name]
                    text += (
                        f"[INVENTARIO FORESTAL] Total de individuos arbóreos "
                        f"contados en esta hoja (conteo de filas de datos, "
                        f"NO el valor máximo de la columna ID): {info['individuos']}\n"
                    )
                    if info.get("volumen_m3") is not None:
                        text += (
                            f"[INVENTARIO FORESTAL] Volumen total (m3) sumado "
                            f"de la columna VT: {info['volumen_m3']}\n"
                        )
                text += df.to_string(index=False) + "\n"

    except Exception as e:
        print(f"Error extrayendo {file_path}: {e}")

    return text


def extract_docx(file_path: str) -> str:
    """
    Extrae texto de un .docx de manera más inteligente:
    - Párrafos normales
    - Tablas: convierte cada fila en "cabecera: valor" para facilitar el regex
    - Preserva contexto de encabezados para que el regex sepa en qué sección está
    """
    doc = Document(file_path)
    parts = []

    for block in iter_blocks(doc):
        if block["type"] == "paragraph":
            txt = block["text"].strip()
            if txt:
                parts.append(txt)
        elif block["type"] == "table":
            parts.append(_table_to_text(block["table"]))

    return "\n".join(parts)


def iter_blocks(doc):
    """
    Itera párrafos y tablas del documento en el orden real del XML.
    python-docx expone doc.paragraphs y doc.tables por separado,
    pero aquí los entregamos en el orden que aparecen en el documento.
    """
    from docx.oxml.ns import qn
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            yield {"type": "paragraph", "text": Paragraph(child, doc).text}

        elif tag == "tbl":
            from docx.table import Table
            yield {"type": "table", "table": Table(child, doc)}


def _table_to_text(table) -> str:
    """
    Convierte una tabla en texto plano preservando contexto.
    Si la primera fila parece ser encabezado, produce líneas como:
    'Nombre especie: Cedrela odorata | Individuos: 8 | Volumen: 0,816 m3'
    Si no tiene encabezado claro, produce líneas separadas por |.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # Eliminar celdas repetidas (python-docx repite celdas combinadas)
        seen = []
        for c in cells:
            if not seen or c != seen[-1]:
                seen.append(c)
        rows.append(seen)

    if not rows:
        return ""

    lines = []
    # Si primera fila tiene texto en todas las celdas → encabezado
    header = rows[0] if all(h for h in rows[0]) else None

    for i, row in enumerate(rows):
        if header and i == 0:
            lines.append(" | ".join(row))
            continue

        if header:
            pairs = []
            for h, v in zip(header, row):
                if v:
                    pairs.append(f"{h}: {v}")
            if pairs:
                lines.append(" | ".join(pairs))
        else:
            lines.append(" | ".join(c for c in row if c))

    return "\n".join(lines)

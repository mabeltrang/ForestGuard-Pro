"""
fecha_checker.py — Detección de la fecha del encabezado de un documento y
evaluación de qué tan desactualizada está frente a hoy.

Busca la fecha en:
- El encabezado real de Word (section.header) para .docx
- Las primeras líneas del cuerpo (muchos formatos oficiales colombianos no
  usan el encabezado de Word como tal, sino que escriben la fecha en las
  primeras líneas, ej. "Bogotá D.C., 15 de enero de 2024")
- Las primeras líneas de la página 1 para .pdf y .xlsx

No asume timezone; compara solo por fecha calendario contra la fecha actual
del sistema.
"""

import re
from datetime import datetime
from typing import Optional

from docx import Document
from pypdf import PdfReader

MESES = {
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
    "julio": 7, "agosto": 8, "septiembre": 9, "setiembre": 9, "octubre": 10,
    "noviembre": 11, "diciembre": 12,
}

_PATRON_TEXTUAL = re.compile(
    r"(\d{1,2})\s+de\s+(" + "|".join(MESES.keys()) + r")\s+(?:de\s+)?(\d{4})",
    re.IGNORECASE,
)
_PATRON_NUMERICO_DMY = re.compile(r"\b(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})\b")
_PATRON_ISO = re.compile(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b")


def _parsear_fechas(texto: str) -> list[datetime]:
    """Encuentra todas las fechas reconocibles en un texto y las parsea."""
    fechas = []

    for m in _PATRON_TEXTUAL.finditer(texto):
        dia, mes_nombre, anio = m.groups()
        mes = MESES.get(mes_nombre.lower())
        try:
            fechas.append(datetime(int(anio), mes, int(dia)))
        except (ValueError, TypeError):
            continue

    for m in _PATRON_ISO.finditer(texto):
        anio, mes, dia = m.groups()
        try:
            fechas.append(datetime(int(anio), int(mes), int(dia)))
        except ValueError:
            continue

    for m in _PATRON_NUMERICO_DMY.finditer(texto):
        a, b, anio = m.groups()
        anio_i = int(anio)
        if anio_i < 100:
            anio_i += 2000 if anio_i < 70 else 1900
        # Formato colombiano estándar es día/mes/año; si el primer número no
        # puede ser mes (>12), confirma día/mes; si ambos son <=12 queda
        # ambiguo y se asume día/mes por convención local.
        dia, mes = int(a), int(b)
        if mes > 12 and dia <= 12:
            dia, mes = mes, dia
        try:
            fechas.append(datetime(anio_i, mes, dia))
        except ValueError:
            continue

    return fechas


def _fecha_mas_reciente_valida(fechas: list[datetime], limite_futuro: Optional[datetime] = None) -> Optional[datetime]:
    """
    De varias fechas candidatas en un mismo bloque de texto (encabezado/pie
    suelen tener solo una, pero por seguridad se filtra), descarta fechas
    absurdamente futuras (probable error de OCR/formato) y retorna la más
    reciente restante como "fecha del documento".
    """
    if not fechas:
        return None
    limite = limite_futuro or datetime.now()
    validas = [f for f in fechas if f.year >= 1990 and f <= limite]
    if not validas:
        # Si todas quedaron en el futuro, no descartar de plano: puede ser
        # una fecha de vigencia/vencimiento, no de expedición. Se retorna la
        # más antigua de las encontradas como mejor esfuerzo.
        return min(fechas)
    return max(validas)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def extraer_fecha_docx(file_path: str, lineas_cuerpo: int = 12) -> Optional[datetime]:
    """
    Busca primero en el encabezado/pie real de Word; si no hay fecha ahí,
    busca en las primeras `lineas_cuerpo` líneas del cuerpo del documento
    (cubre el caso, muy común en oficios colombianos, de escribir la fecha
    como texto normal al inicio en vez de usar el encabezado de Word).
    """
    doc = Document(file_path)

    texto_encabezado = ""
    for section in doc.sections:
        for p in section.header.paragraphs:
            texto_encabezado += p.text + "\n"
        for p in section.footer.paragraphs:
            texto_encabezado += p.text + "\n"

    fechas = _parsear_fechas(texto_encabezado)
    fecha = _fecha_mas_reciente_valida(fechas)
    if fecha:
        return fecha

    texto_cuerpo = "\n".join(p.text for p in doc.paragraphs[:lineas_cuerpo])
    fechas = _parsear_fechas(texto_cuerpo)
    return _fecha_mas_reciente_valida(fechas)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extraer_fecha_pdf(file_path: str, caracteres_inicio: int = 700) -> Optional[datetime]:
    """
    Busca la fecha en los primeros `caracteres_inicio` caracteres de texto
    extraído de la página 1 (aproxima "el encabezado" ya que pypdf no expone
    posición de layout).
    """
    try:
        reader = PdfReader(file_path)
        if not reader.pages:
            return None
        texto = reader.pages[0].extract_text() or ""
    except Exception:
        return None

    fechas = _parsear_fechas(texto[:caracteres_inicio])
    fecha = _fecha_mas_reciente_valida(fechas)
    if fecha:
        return fecha
    # Si no apareció en el bloque inicial, buscar en toda la página 1 como
    # respaldo (algunos PDF traen la fecha más abajo, ej. firma/radicado)
    fechas = _parsear_fechas(texto)
    return _fecha_mas_reciente_valida(fechas)


# ---------------------------------------------------------------------------
# PUNTO DE ENTRADA ÚNICO + EVALUACIÓN DE VIGENCIA
# ---------------------------------------------------------------------------

def extraer_fecha_documento(file_path: str, nombre_archivo: str = "") -> Optional[datetime]:
    nombre = nombre_archivo or file_path
    ext = nombre.lower().rsplit(".", 1)[-1] if "." in nombre else ""
    try:
        if ext in ("docx", "doc"):
            return extraer_fecha_docx(file_path)
        elif ext == "pdf":
            return extraer_fecha_pdf(file_path)
    except Exception:
        return None
    return None


def evaluar_vigencia(
    fecha: Optional[datetime],
    umbral_advertencia_dias: int = 180,
    umbral_desactualizado_dias: int = 365,
    referencia: Optional[datetime] = None,
) -> dict:
    """
    Clasifica qué tan vieja es la fecha encontrada frente a `referencia`
    (hoy, por defecto). Los umbrales son un punto de partida razonable para
    trámites ambientales (una resolución/certificado con más de un año suele
    necesitar verificarse si sigue vigente); ajústalos si tu trámite
    específico tiene un plazo de vigencia distinto.

    Retorna {"estado": "vigente"|"revisar"|"desactualizado"|"sin_fecha",
             "dias": int|None, "fecha": datetime|None}.
    """
    ref = referencia or datetime.now()
    if fecha is None:
        return {"estado": "sin_fecha", "dias": None, "fecha": None}

    dias = (ref - fecha).days

    if dias < 0:
        estado = "fecha_futura"
    elif dias > umbral_desactualizado_dias:
        estado = "desactualizado"
    elif dias > umbral_advertencia_dias:
        estado = "revisar"
    else:
        estado = "vigente"

    return {"estado": estado, "dias": dias, "fecha": fecha}

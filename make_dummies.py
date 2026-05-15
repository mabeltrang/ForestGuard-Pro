from docx import Document

d1 = Document()
d1.add_paragraph("Reporte Técnico Forestal 1\n\nResumen:\nLa Zona de Vida identificada fue: bosque húmedo tropical.\nEl área estudiada corresponde a 150 hectáreas.\nSe contabilizaron en total 543 individuos.\nEl costo total estimado fue de $5,000,000.\nEl volumen maderable es de 45.2 m3.")
d1.save("doc1.docx")

d2 = Document()
d2.add_paragraph("Reporte Técnico Forestal 2\n\nResumen de Campo:\nSe detectó que la zona de vida es bosque seco tropical.\nEl área intervenida es 150 hectáreas.\nIndividuos muestreados: 543.\nPresupuesto: $5,000,000.\nVolumen total: 45.2 m³.")
d2.save("doc2.docx")

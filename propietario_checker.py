# -*- coding: utf-8 -*-
"""
propietario_checker.py — Verifica que el PROPIETARIO reportado en el
Inventario Forestal (xlsx) coincida con el propietario real identificado
en el Plan/Informe de Aprovechamiento Forestal (docx).

Contexto: el "PROPIETARIO" del predio nunca es Unergy (ni Solenium, ni
ningún otro contratista/solicitante) — es la persona natural o jurídica
dueña del predio, la que autoriza el aprovechamiento. El informe de
aprovechamiento siempre menciona su nombre completo junto a su cédula/NIT
(ej. "... se cuenta con la autorización de la propietaria, la señora
Ruby Esther Difilippo Cadena identificada con C.C. 43.081.891 ..."), así
que ese es el nombre contra el cual se debe validar el inventario.
"""

import difflib
import re
import unicodedata
from typing import Optional

import openpyxl
from docx import Document

_PATRON_PROPIETARIO = re.compile(
    r"propietari[ao]s?\b.{0,60}?"
    r"([A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÑáéíóúñ0-9&.\- ]{3,80}?)"
    r"\s*,?\s*identificad[oa]?\s+con\s+"
    r"(C\.?\s?C\.?|NIT|C\.?E\.?|Ced(?:ula)?)\.?\s*([\d.,\-]+)",
    re.IGNORECASE,
)

_PREFIJOS_TRATAMIENTO = re.compile(
    r"^(la\s+se[ñn]ora|el\s+se[ñn]or(?:\(a\))?|la\s+sociedad|el\s+se[ñn]or\s*a)\s+",
    re.IGNORECASE,
)

_ETIQUETAS_PROPIETARIO_XLSX = re.compile(r"(?i)^propietari[ao]s?(\s*\(?s?\)?)?\s*:?$")


# ---------------------------------------------------------------------------
# EXTRACCIÓN — Plan/Informe de Aprovechamiento Forestal (.docx)
# ---------------------------------------------------------------------------

def extraer_propietarios_docx(file_path: str) -> list:
    """
    Busca todas las menciones "propietario/a ... identificado con C.C./NIT
    ..." en párrafos y tablas del documento. Retorna una lista (puede haber
    más de un propietario, ej. predios colindantes de distintos dueños):
    [{"nombre": str, "tipo_id": str, "numero_id": str}, ...]
    """
    doc = Document(file_path)
    textos = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                textos.append(cell.text)

    texto_completo = "\n".join(textos)
    hallazgos = []
    vistos = set()
    for m in _PATRON_PROPIETARIO.finditer(texto_completo):
        nombre_raw, tipo_id, numero_id = m.groups()
        nombre = _PREFIJOS_TRATAMIENTO.sub("", nombre_raw).strip(" ,.")
        clave = (nombre.upper(), numero_id)
        if clave in vistos or not nombre:
            continue
        vistos.add(clave)
        hallazgos.append({
            "nombre": nombre,
            "tipo_id": tipo_id.replace(" ", "").upper().rstrip("."),
            "numero_id": numero_id,
        })
    return hallazgos


# ---------------------------------------------------------------------------
# EXTRACCIÓN — Inventario Forestal (.xlsx)
# ---------------------------------------------------------------------------

def extraer_propietario_xlsx(file_path: str, max_filas: int = 20, max_cols: int = 20) -> Optional[str]:
    """
    Busca en las primeras `max_filas` x `max_cols` celdas de cada hoja una
    etiqueta tipo "PROPIETARIO" / "PROPIETARIO(S)" y retorna el valor de la
    primera celda no vacía a su derecha en la misma fila (soporta celdas
    combinadas, ya que openpyxl solo expone el valor en la celda superior
    izquierda del rango combinado).
    """
    wb = openpyxl.load_workbook(file_path, data_only=True)
    for ws in wb.worksheets:
        max_r = min(max_filas, ws.max_row or 0)
        max_c = min(max_cols, ws.max_column or 0)
        for r in range(1, max_r + 1):
            for c in range(1, max_c + 1):
                valor = ws.cell(row=r, column=c).value
                if not isinstance(valor, str):
                    continue
                if _ETIQUETAS_PROPIETARIO_XLSX.match(valor.strip()):
                    for c2 in range(c + 1, max_c + 1):
                        candidato = ws.cell(row=r, column=c2).value
                        if candidato not in (None, ""):
                            return str(candidato).strip()
    return None


# ---------------------------------------------------------------------------
# COMPARACIÓN
# ---------------------------------------------------------------------------

def _normalizar(texto: str) -> str:
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode()
    texto = re.sub(r"[^A-Za-z0-9 ]", " ", texto)
    return " ".join(texto.upper().split())


def _similitud(a: str, b: str) -> float:
    a_norm, b_norm = _normalizar(a), _normalizar(b)
    if not a_norm or not b_norm:
        return 0.0
    ratio = difflib.SequenceMatcher(None, a_norm, b_norm).ratio()
    if a_norm in b_norm or b_norm in a_norm:
        ratio = max(ratio, 0.9)
    return ratio


def comparar_propietario(valor_inventario: Optional[str], propietarios_docx: list, umbral: float = 0.6) -> dict:
    """
    Compara el valor del inventario contra cada propietario extraído del
    plan de aprovechamiento y retorna la mejor coincidencia.

    {"coincide": bool|None, "mejor_candidato": str|None, "similitud": float|None}
    """
    if not valor_inventario or not propietarios_docx:
        return {"coincide": None, "mejor_candidato": None, "similitud": None}

    mejor_nombre = None
    mejor_ratio = 0.0
    for p in propietarios_docx:
        ratio = _similitud(valor_inventario, p["nombre"])
        if ratio > mejor_ratio:
            mejor_ratio = ratio
            mejor_nombre = p["nombre"]

    return {
        "coincide": mejor_ratio >= umbral,
        "mejor_candidato": mejor_nombre,
        "similitud": round(mejor_ratio, 2),
    }


# ---------------------------------------------------------------------------
# PUNTO DE ENTRADA ÚNICO
# ---------------------------------------------------------------------------

def verificar_propietario_inventario(path_inventario_xlsx: str, path_plan_docx: str) -> dict:
    """
    {
        "propietario_inventario": str|None,
        "propietarios_plan": [{"nombre","tipo_id","numero_id"}, ...],
        "coincide": bool|None,   # None = no se pudo determinar (falta algún dato)
        "mejor_candidato": str|None,
        "similitud": float|None,
    }
    """
    propietario_inventario = extraer_propietario_xlsx(path_inventario_xlsx)
    propietarios_plan = extraer_propietarios_docx(path_plan_docx)
    comparacion = comparar_propietario(propietario_inventario, propietarios_plan)

    return {
        "propietario_inventario": propietario_inventario,
        "propietarios_plan": propietarios_plan,
        **comparacion,
    }

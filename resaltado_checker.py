"""
resaltado_checker.py — Detección de resaltado AMARILLO en documentos.

En Unergy el resaltado amarillo se usa como convención interna para marcar
texto o zonas de una imagen/plano donde falta un dato antes de radicar. Este
módulo detecta:

1. Texto resaltado en amarillo dentro de párrafos y tablas de un .docx
   (incluyendo cuadros de texto flotantes), leyendo el atributo nativo de
   resaltado de Word (w:highlight val="yellow"), NO un color de fuente.
2. Imágenes embebidas (docx) o páginas rasterizadas (pdf) que contienen un
   BLOQUE contiguo de amarillo tipo marcador (no solo unos pocos píxeles
   sueltos de un logo o de terreno en un mapa — se exige un blob conectado
   de tamaño mínimo para contar como resaltado real).
3. Anotaciones de tipo "Highlight" nativas de PDF.

Cada hallazgo se reporta como SÍ/NO (no como porcentaje) y, cuando es
posible, con la SECCIÓN del documento en la que aparece:
- En .docx: el título (estilo "Heading"/"Título") más reciente antes del
  hallazgo, tal como se ve en el panel de Navegación de Word.
- En .pdf: la entrada de la tabla de contenido/marcadores (outline) que
  cubre esa página, si el PDF tiene marcadores.

No requiere API externa: todo es análisis local (python-docx / pypdf /
PyMuPDF + Pillow + scipy). Si PyMuPDF o scipy no están instalados, esas
partes del análisis se omiten silenciosamente.
"""

import io
import re
from typing import Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from pypdf import PdfReader

try:
    from PIL import Image
    import numpy as np
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

try:
    from scipy import ndimage
    _SCIPY_OK = True
except ImportError:
    _SCIPY_OK = False

try:
    import fitz  # PyMuPDF
    _FITZ_OK = True
except ImportError:
    _FITZ_OK = False


_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_MC_CHOICE = f"{{{_MC_NS}}}Choice"
_MC_FALLBACK = f"{{{_MC_NS}}}Fallback"

_PATRON_TITULO_ESTILO = re.compile(r"(?i)^(heading|t[íi]tulo)\s*\d*$")


def _es_estilo_titulo(paragraph) -> bool:
    try:
        nombre = paragraph.style.name or ""
    except Exception:
        return False
    return bool(_PATRON_TITULO_ESTILO.match(nombre.strip()))


# ---------------------------------------------------------------------------
# UTILIDAD: ¿esta imagen tiene un bloque contiguo de amarillo tipo marcador?
# ---------------------------------------------------------------------------

def _tiene_bloque_amarillo(imagen_bytes: bytes, area_min_pct: float = 2.5) -> bool:
    """
    True si la imagen contiene un blob CONTIGUO de amarillo tipo resaltador
    (H≈45-68°, saturación y brillo altos) que cubre al menos `area_min_pct`%
    del área de la imagen. Esto descarta ruido disperso (antialiasing de un
    logo, un pixel amarillo en un mapa, etc.) que no forma un bloque real.
    """
    if not _PIL_OK:
        return False
    try:
        img = Image.open(io.BytesIO(imagen_bytes)).convert("RGB")
        img.thumbnail((400, 400))
        arr = np.asarray(img).astype("float32") / 255.0
        r, g, b = arr[..., 0], arr[..., 1], arr[..., 2]

        maxc = np.max(arr, axis=-1)
        minc = np.min(arr, axis=-1)
        delta = maxc - minc

        hue = np.zeros_like(maxc)
        mask_d = delta > 1e-6
        idx_r = mask_d & (maxc == r)
        idx_g = mask_d & (maxc == g) & ~idx_r
        with np.errstate(divide="ignore", invalid="ignore"):
            hue[idx_r] = (60 * ((g[idx_r] - b[idx_r]) / delta[idx_r]) + 360) % 360
            hue[idx_g] = 60 * ((b[idx_g] - r[idx_g]) / delta[idx_g]) + 120

        sat = np.where(maxc > 0, delta / np.maximum(maxc, 1e-6), 0)
        val = maxc

        amarillo = (hue >= 45) & (hue <= 68) & (sat >= 0.35) & (val >= 0.55)
        if not amarillo.any():
            return False

        if not _SCIPY_OK:
            # Sin scipy, usar el % total como respaldo con umbral más alto
            return float(100.0 * amarillo.sum() / amarillo.size) >= (area_min_pct * 2)

        etiquetas, n = ndimage.label(amarillo)
        if n == 0:
            return False
        tamanos = ndimage.sum(amarillo, etiquetas, range(1, n + 1))
        pct_mayor_blob = 100.0 * tamanos.max() / amarillo.size
        return pct_mayor_blob >= area_min_pct
    except Exception:
        return False


# ---------------------------------------------------------------------------
# DOCX — texto e imágenes resaltadas, con sección
# ---------------------------------------------------------------------------

def _runs_resaltados_amarillo(paragraph) -> list:
    """Retorna los tramos de texto de un párrafo marcados en amarillo."""
    hallazgos = []
    fragmento = ""
    for run in paragraph.runs:
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW and run.text.strip():
            fragmento += run.text
        else:
            if fragmento.strip():
                hallazgos.append(fragmento.strip())
            fragmento = ""
    if fragmento.strip():
        hallazgos.append(fragmento.strip())
    return hallazgos


def _extraer_blip_rids(element) -> list:
    """
    Retorna los r:embed de las imágenes (a:blip) dentro de `element`,
    excluyendo las que están en <mc:Fallback> (copia legada VML duplicada
    de un cuadro de texto/forma, para no contar la misma imagen dos veces).
    """
    rids = []
    for blip in element.iter(qn("a:blip")):
        en_fallback = any(anc.tag == _MC_FALLBACK for anc in blip.iterancestors())
        if en_fallback:
            continue
        rid = blip.get(qn("r:embed"))
        if rid:
            rids.append(rid)
    return rids


def _resolver_imagen(part, rid) -> Optional[bytes]:
    try:
        return part.rels[rid].target_part.blob
    except Exception:
        return None


def _procesar_parrafo(paragraph, part, seccion: Optional[str], tipo_bloque: str,
                       hallazgos_texto: list, hallazgos_img: list, area_min_pct: float) -> None:
    for texto in _runs_resaltados_amarillo(paragraph):
        hallazgos_texto.append({
            "texto_resaltado": texto,
            "seccion": seccion,
            "ubicacion": tipo_bloque,
        })
    for rid in _extraer_blip_rids(paragraph._p):
        data = _resolver_imagen(part, rid)
        if data and _tiene_bloque_amarillo(data, area_min_pct):
            hallazgos_img.append({
                "seccion": seccion,
                "ubicacion": tipo_bloque,
            })


def _procesar_tabla(table, part, seccion: Optional[str], hallazgos_texto: list,
                     hallazgos_img: list, area_min_pct: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _procesar_parrafo(p, part, seccion, "tabla", hallazgos_texto, hallazgos_img, area_min_pct)


def _procesar_textboxes(element, doc, part, seccion: Optional[str], hallazgos_texto: list,
                         hallazgos_img: list, area_min_pct: float) -> None:
    """Recorre cuadros de texto (<w:txbxContent>), solo la copia mc:Choice."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    def _revisar(txbx):
        for sub in txbx:
            subtag = sub.tag.split("}")[-1] if "}" in sub.tag else sub.tag
            if subtag == "p":
                _procesar_parrafo(Paragraph(sub, doc), part, seccion, "cuadro de texto",
                                   hallazgos_texto, hallazgos_img, area_min_pct)
            elif subtag == "tbl":
                _procesar_tabla(Table(sub, doc), part, seccion, hallazgos_texto, hallazgos_img, area_min_pct)

    choices = list(element.iter(_MC_CHOICE))
    if choices:
        for choice in choices:
            for txbx in choice.iter(qn("w:txbxContent")):
                _revisar(txbx)
    else:
        for txbx in element.iter(qn("w:txbxContent")):
            _revisar(txbx)


def detectar_resaltado_docx(file_path: str, area_min_pct: float = 2.5) -> dict:
    """
    Recorre el documento en orden real, siguiendo el título (estilo
    Heading/Título) vigente en cada punto, y retorna:
    {"texto": [{"texto_resaltado","seccion","ubicacion"}, ...],
     "imagenes": [{"seccion","ubicacion"}, ...]}
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(file_path)
    part = doc.part
    hallazgos_texto: list = []
    hallazgos_img: list = []
    seccion_actual: Optional[str] = None

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            p = Paragraph(child, doc)
            if _es_estilo_titulo(p) and p.text.strip():
                seccion_actual = p.text.strip()
            _procesar_parrafo(p, part, seccion_actual, "párrafo", hallazgos_texto, hallazgos_img, area_min_pct)
            _procesar_textboxes(child, doc, part, seccion_actual, hallazgos_texto, hallazgos_img, area_min_pct)

        elif tag == "tbl":
            _procesar_tabla(Table(child, doc), part, seccion_actual, hallazgos_texto, hallazgos_img, area_min_pct)
            _procesar_textboxes(child, doc, part, seccion_actual, hallazgos_texto, hallazgos_img, area_min_pct)

    # Encabezados y pies de página (no llevan "sección" propia del cuerpo)
    for section in doc.sections:
        for contenedor, etiqueta in ((section.header, "encabezado"), (section.footer, "pie de página")):
            hpart = contenedor.part
            for p in contenedor.paragraphs:
                _procesar_parrafo(p, hpart, None, etiqueta, hallazgos_texto, hallazgos_img, area_min_pct)
            for table in contenedor.tables:
                _procesar_tabla(table, hpart, None, hallazgos_texto, hallazgos_img, area_min_pct)

    return {"texto": hallazgos_texto, "imagenes": hallazgos_img}


# ---------------------------------------------------------------------------
# PDF — mapa de secciones a partir de marcadores/outline
# ---------------------------------------------------------------------------

def _mapa_secciones_pdf(reader: PdfReader) -> list:
    """Lista [(indice_pagina, titulo), ...] ordenada, a partir del outline del PDF."""
    entradas = []

    def _recorrer(items):
        for item in items:
            if isinstance(item, list):
                _recorrer(item)
                continue
            try:
                pagina = reader.get_destination_page_number(item)
                entradas.append((pagina, str(item.title)))
            except Exception:
                continue

    try:
        _recorrer(reader.outline)
    except Exception:
        pass
    entradas.sort(key=lambda x: x[0])
    return entradas


def _seccion_para_pagina(mapa: list, pagina_idx: int) -> Optional[str]:
    seccion = None
    for pidx, titulo in mapa:
        if pidx <= pagina_idx:
            seccion = titulo
        else:
            break
    return seccion


# ---------------------------------------------------------------------------
# PDF — anotaciones nativas de resaltado
# ---------------------------------------------------------------------------

def detectar_resaltado_anotaciones_pdf(file_path: str) -> list:
    hallazgos = []
    try:
        reader = PdfReader(file_path)
        mapa = _mapa_secciones_pdf(reader)
        for i, page in enumerate(reader.pages):
            annots = page.get("/Annots")
            if not annots:
                continue
            for ref in annots:
                try:
                    obj = ref.get_object()
                except Exception:
                    continue
                if obj.get("/Subtype") != "/Highlight":
                    continue
                color = obj.get("/C")
                es_amarillo = True
                if color and len(color) == 3:
                    r, g, b = [float(c) for c in color]
                    es_amarillo = r > 0.7 and g > 0.7 and b < 0.5
                if es_amarillo:
                    hallazgos.append({
                        "pagina": i + 1,
                        "seccion": _seccion_para_pagina(mapa, i),
                        "comentario": str(obj.get("/Contents")) if obj.get("/Contents") else None,
                    })
    except Exception:
        pass
    return hallazgos


# ---------------------------------------------------------------------------
# PDF — bloques amarillos en páginas rasterizadas
# ---------------------------------------------------------------------------

def detectar_resaltado_imagenes_pdf(file_path: str, area_min_pct: float = 2.5, max_paginas: int = 60) -> list:
    hallazgos = []
    if not _FITZ_OK or not _PIL_OK:
        return hallazgos
    try:
        reader = PdfReader(file_path)
        mapa = _mapa_secciones_pdf(reader)
    except Exception:
        mapa = []
    try:
        doc = fitz.open(file_path)
        for i in range(min(len(doc), max_paginas)):
            page = doc[i]
            mat = fitz.Matrix(100 / 72, 100 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            if _tiene_bloque_amarillo(pix.tobytes("png"), area_min_pct):
                hallazgos.append({"pagina": i + 1, "seccion": _seccion_para_pagina(mapa, i)})
        doc.close()
    except Exception:
        pass
    return hallazgos


# ---------------------------------------------------------------------------
# PUNTO DE ENTRADA ÚNICO
# ---------------------------------------------------------------------------

def detectar_resaltado_documento(file_path: str, nombre_archivo: str = "") -> dict:
    """
    Dispatcher por extensión. Retorna:
    {
        "texto": [{"texto_resaltado","seccion","ubicacion"}, ...],
        "imagenes": [{"seccion","ubicacion"} o {"pagina","seccion"}, ...],
        "anotaciones_pdf": [{"pagina","seccion","comentario"}, ...],
        "tiene_hallazgos": bool,
    }
    """
    nombre = nombre_archivo or file_path
    ext = nombre.lower().rsplit(".", 1)[-1] if "." in nombre else ""

    resultado = {"texto": [], "imagenes": [], "anotaciones_pdf": [], "tiene_hallazgos": False}

    if ext in ("docx", "doc"):
        try:
            r = detectar_resaltado_docx(file_path)
            resultado["texto"] = r["texto"]
            resultado["imagenes"] = r["imagenes"]
        except Exception:
            pass

    elif ext == "pdf":
        resultado["anotaciones_pdf"] = detectar_resaltado_anotaciones_pdf(file_path)
        resultado["imagenes"] = detectar_resaltado_imagenes_pdf(file_path)

    resultado["tiene_hallazgos"] = bool(
        resultado["texto"] or resultado["imagenes"] or resultado["anotaciones_pdf"]
    )
    return resultado
